from os import chdir, getcwd
import wikipedia
from docx.api import Document


def main():   #* defining the main function for the program

    page = wikipedia.random(pages = 1)          #* Fetching a random wikipedia page title
    print("The title of the page is: ", page)   #* Displaying the title to user
    ask = input("Do you want to fetch this page?[Y/N] >>> ").lower();print()    #* asking for confirmation
    
    try:       #* Try block
        if ask == 'y':        #* if user confirms to get the article

            result = wikipedia.page(page)  #* fetch the page as page object
            doc = Document()         #* make the document object to save the fetched results

            def choice():   #* Fetching the article as summary/whole and asking user for another article
            
                def again():
                    again = input("Do you want another article?[Y/N] >>> ").lower()  #* asking user for another article
            
                    if again == 'y':main()    #* fetch a new wikipedia page title 
            
                    elif again == 'n':exit()  #* Exit the program
            
                    else:print("Invalid choice!!");exit() 
                            
                choose = int(input("Do you want to get summary or Whole Article?\nChoose 1 for summary & 2 for Whole Article: "))
                #* Giving a choice to choose between summary/Whole article
                
                if choose == 1:                 #* User choose to get summary
                    summ = result.summary       #* fetch summary for the page result
                    doc.add_heading(page)       #* add the heading(same as the title of the page fetched)
                    doc.add_paragraph(summ)     #* Add the fetched summary to the document
                    
                    change_dir()                #* ask the user to change the directory of the save file

                    doc.save(f"{page}_Summary.docx")    #* Saving the file as the title + 'summary' (to indicate summary doc) 
                    print(f"Document Saved at {getcwd()}")   #* Give confirmation of document saved
                    #* Getcwd will give user the location of saved doc 
                    
                    again()                     #* if they want a new article or not     

                elif choose == 2:               #* User chooses to fetch the whole article
                    content = result.content    #* fetching the content from the page
                    doc.add_heading(page)       #* add the heading(same as the title of the page fetched)
                    doc.add_paragraph(content)  #* add the fetched content to the document
                    
                    change_dir() #* ask the user to change the directory of the save file

                    doc.save(f"{page}.docx")    #* Saving the file as the 'title'.docx
                    print(f"Document Saved at {getcwd()}") #* gives confirmation of document saved
                    #* Getcwd will give user the location of saved doc 

                    again()                     #* if they want a new article or not

                else:print("Invalid Choice!");choice()   #* Ask the user again to get summary or whole article if wrong input
            
            choice()   #* Run the choice function
        
        elif ask == 'n':      #* if user did not want the article
            print("Finding a new article..")    
            main()           #* A new article will be fetched
        
        else:print("Invalid choice");exit()                        #* If user gices wrong choice on ask variable
    
    except wikipedia.exceptions.PageError:    #* handling the error (Wikipedia article not found)
        print("Sorry but the page was not found..")
        print("Searching another page for you..")
        main()                                #* Finding a new random article 

def change_dir():   #* Helping user to change the directory of the save file
    
    def direct():   #* User inputs the path of the directory they want to save to 
        direc = input("Enter the directory path you want to save the file to: ")
        chdir(direc)

    def direct_again():
        try:
            direct()   #* Another try at specifi=ying the location
        except FileNotFoundError:    #* If the save path isnt found after the second try, doc saves at default path
            print("Enough tries..Saving at defualt path instead")
    
    ask = int(input(f"Do you want to save at default or specified location?\n1. Default- '{getcwd()}' \n2. Give save path\n\
Your choice: "))          #* Asking the user to specify the save location
    
    try:

        if ask == 1:pass            #* Default location (Will save in the current working directory)
        elif ask==2:direct()        #* if they want doc at a specific path 
        else:print("invalid option..try again");direct()    #* On invalid option

    except FileNotFoundError:       #* Error handling (If the save path specified was not found)

        print("No such location found...")     
        choice = input("Save at defualt path instead?[Y/N] >>>").lower()   #* Will prompt user to save at default path

        if choice=='y':pass     #* If the user couldnt specify a location and wants the file at default path instead
        
        elif choice=='n':direct_again()      #* Another try at specifying the location of the saved document
        
        else:                       #* if still the user couldn't give the correct path, doc will save at the default path 
            print("Invalid option")
            print("Saving at default path instead..")
            pass
try:
    main()     #* Calling the main function to execute the program
except Exception:      #* Error Handling
    print("Please check your internet connection before trying again.")